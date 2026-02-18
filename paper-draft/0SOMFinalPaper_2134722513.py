import itertools
from pathlib import Path

import minisom
import numpy as np
import pandas as pd
import plotly.express as px
import plotly.graph_objects as go
import plotly.io as pio
import shapely
import matplotlib.pyplot as plt

from plotly.colors import sample_colorscale, DEFAULT_PLOTLY_COLORS
from plotly.subplots import make_subplots

from scipy.cluster import hierarchy
from sklearn.metrics import silhouette_score
from sklearn.decomposition import PCA
from IPython.display import Image
from scipy.interpolate import griddata
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH


# Parameters

data_path = Path(__file__).parent / 'a_data_Paper.csv'
input_columns = ['F4','F5','F7','F8','F9', 'pH_L']



som_nx = None  # Set to None to determine these automatically (PCA)
som_ny = None
#________________________________________________



# Exp 3 clusters with QE = 0.13 and TE = 0.0
som_sigma = 2.4
som_learning_rate = 0.37
som_random_seed = 10
som_max_iterations = 6000



som_activation_distance = 'euclidean'
som_topology = 'hexagonal'
som_neighborhood_function = 'gaussian'


# Paths

links_data = Path(__file__).parent / "links.txt"

max_clusters = 10  # Maximum number of clusters to test
 
number_of_clusters = None  # Set to None to pick the optimal number automatically



# Visualization
#features_colorscale = 'spectral'
features_colorscale = 'RdBu'
pio.templates.default = 'plotly_white'
pio.renderers.default = 'browser' # Set to browser if not already
plotly_interactive = True

# Convenience function to allow changing between interactive and static images
def show_figure(fig: go.Figure):
    if plotly_interactive:
        fig.show()
    else:
        return Image(data=fig.to_image(format='png', scale=2))


# Read input data
df = pd.read_csv(data_path)


# Apply data transformation

# Normalization:
def apply_normalization(arr):
    data_min = arr.min(axis=0)
    data_range = arr.max(axis=0) - data_min

    return (arr - data_min) / data_range

# Transforms data into a range [-1, 1].
    #return ((2*(arr - data_min))/ data_range)-1


# # Standardization
# def apply_standardization(arr):
#     data_mean = arr.mean(axis=0)
#     data_std = arr.std(axis=0)
#     return (arr - data_mean) / data_std

# # Sigmoid
# def apply_sigmoid(arr):
#     return 1 / (1 + np.exp(-arr))


input_array = df[input_columns].values

## Uncomment the appropriate transform
transformed_data = apply_normalization(input_array)
# transformed_data = apply_standardization(input_array)
# transformed_data = apply_sigmoid(input_array)

#print(f"Transformed data shape: {transformed_data}")

# Derive SOM dimensions
if som_nx is None or som_ny is None:
    recommended_nodes = 5 * np.sqrt(len(df))
    print(f"Recommended number of nodes: {recommended_nodes:.0f}")
   
    # "Naive" approach: square root of the number of samples (with round-off)
    # final_nx = np.round(np.sqrt(som_m)).astype(int)
    # final_ny = np.round(som_m / som_nx).astype(int)

    # Perform PCA analysis
    pca = PCA(n_components=2)
    pca.fit(transformed_data)

    # Calculate the ratio between the first two principal components
    pca_ratio = pca.explained_variance_[0] / pca.explained_variance_[1]
    final_nx = np.round(np.sqrt(recommended_nodes * pca_ratio)).astype(int)
    final_ny = np.round(recommended_nodes / final_nx).astype(int)
    print(f"Using PCA to determine SOM dimensions: {final_nx} x {final_ny} (PCA ratio: {pca_ratio:.2f})")
else:
    final_nx, final_ny = som_nx, som_ny
    print(f"Using fixed SOM dimensions: {final_nx} x {final_ny}")

# Print all principal components values
print(f"Principal component values: {pca.explained_variance_}")
#print(f"Principal component ratio: {pca.explained_variance_ratio_}")



# Set up model
som = minisom.MiniSom(
    final_nx,
    final_ny,
    len(input_columns),
    sigma=som_sigma,
    learning_rate=som_learning_rate,
    activation_distance=som_activation_distance,
    topology=som_topology,
    neighborhood_function=som_neighborhood_function,
    random_seed=som_random_seed,
)

# Train
som.train(transformed_data, som_max_iterations)

# Extract data
xx, yy = som.get_euclidean_coordinates()
umatrix = som.distance_map()
weights = som.get_weights()
win_map = som.win_map(transformed_data, return_indices = True)

# Print quantization and topographic errors for the trained model
print(f"Quantization error: {som.quantization_error(transformed_data)}")
print(f"Topographic error : {som.topographic_error(transformed_data)}")


# Range of SOM grid sizes to test
grid_sizes = range(2, 14)

# Store QE and TE for each grid size
qe_values = []
te_values = []

for size in grid_sizes:
    # Initialize and train SOM
    som = minisom.MiniSom(size, size, transformed_data.shape[1], som_sigma, som_learning_rate)
    som.random_weights_init(transformed_data)
    som.train_random(transformed_data, som_max_iterations)

    # Calculate and store QE
    qe = som.quantization_error(transformed_data)
    qe_values.append(qe)

     # Calculate and store TE
    te = som.topographic_error(transformed_data)
    te_values.append(te)
   

# Plotting QE vs SOM size
plt.plot(grid_sizes, qe_values, marker='o')
plt.xlabel("SOM Grid Size")
plt.ylabel("Quantization Error")
plt.title("Quantization Error vs. SOM Size")
plt.xticks(grid_sizes)
plt.grid(True)
#plt.show()


# Plotting TE vs SOM size
plt.plot(grid_sizes, te_values, marker='o')
plt.xlabel("SOM Grid Size")
plt.ylabel("Topographic Error")
plt.title("Topographic Error vs SOM Size")
plt.xticks(grid_sizes)
plt.grid(True)
#plt.show()

#Printing QE and TE values

# for size in grid_sizes:
#     print(f"Quantization error: Som size {qe_values, size}\n")

# for size in grid_sizes:
#     print(f"Topographic error: Som size {te_values, size}\n")

# #Plot U-Matrix
# fig = px.imshow(
#     umatrix,
#     color_continuous_scale=features_colorscale,
#     title="U-Matrix",
#     labels=dict(x="X", y="Y"),
    
# )
#     # print("\n")
#     # print("\t\n")
#     # print("\n")

# Setup input to clustering and dendrogram plot
mij = np.meshgrid(np.arange(final_nx), np.arange(final_ny), indexing='ij')
cells_ij = np.column_stack((mij[0].flat, mij[1].flat))
labels = [(i, j) for i, j in cells_ij]
flat_weights = weights.reshape(-1, len(input_columns)) 

# --------------------------------------------------------
# Clustering and dendrogram
# --------------------------------------------------------

# Calculate links and save to file
links = hierarchy.linkage(flat_weights, method='complete', metric='euclidean')
np.savetxt(links_data, links)

# Clustering

cluster_results = []

for i in range(2, max_clusters):
    flat_labels = hierarchy.fcluster(links, i, 'maxclust')
    #dbi = davies_bouldin_score(flat_weights, flat_labels) if len(set(flat_labels)) == i else None
    sil = silhouette_score(flat_weights, flat_labels) if len(set(flat_labels)) == i else None
    cluster_results.append({
        "Number of clusters": i,
        #"DBI": dbi,
        "Silhouette": sil,
        "labels": flat_labels,
    })

cluster_df = pd.DataFrame.from_dict(cluster_results)


#dbi_optimal_cluster = cluster_df.iloc[cluster_df['DBI'].idxmin()]
sil_optimal_cluster = cluster_df.iloc[cluster_df['Silhouette'].idxmax()]

if number_of_clusters is None:
    selected_cluster = sil_optimal_cluster
    #selected_cluster = dbi_optimal_cluster
else:
    selected_cluster = cluster_df[cluster_df['Number of clusters'] == number_of_clusters].iloc[0]

#print(f"Optimal number of clusters based on DBI        : {dbi_optimal_cluster['Number of clusters']}")
print(f"Optimal number of clusters based on Silhouette : {sil_optimal_cluster['Number of clusters']}")

print()

if number_of_clusters is None:
    print(f"Number of clusters picked based on Silhouette     : {selected_cluster['Number of clusters']}")
    # print(f"Number of clusters picked based on DBI         : {selected_cluster['Number of clusters']}")
else:
    print(f"Number of clusters fixed to                    : {number_of_clusters}")

# fig = px.line(cluster_df, x='Number of clusters', y='DBI', markers=True)
# fig.layout.title = "Davies-Bouldin Index (lower is better)"
# fig.layout.yaxis.spikesnap = "hovered data"
# fig.layout.yaxis.spikemode = "across"
# fig.layout.yaxis.spikethickness = 1

fig = px.line(cluster_df, x='Number of clusters', y='Silhouette', markers=True)
fig.layout.title = "Silhouette score (higher is better)"
fig.layout.yaxis.spikesnap = "hovered data"
fig.layout.yaxis.spikemode = "across"
fig.layout.yaxis.spikethickness = 2
#fig
show_figure(fig)
#plt.show()


# Save separate CSV file the cluster results with silhouette scores
cluster_df.to_excel("Silhouette_cluster_results.xlsx", index=False)

# Optimal number of lables
node_labels = selected_cluster['labels'].reshape(final_nx, final_ny)

# Sample labels
sample_ij = [None] * len(df)
sample_label = [-1] * len(df)
for key, value in win_map.items():
    for ix in value:
        sample_ij[ix] = key
        sample_label[ix] = node_labels[key]

df['ij'] = sample_ij
df['cluster_label'] = sample_label


n_cols = 8
for lab, sub_df in df.groupby('cluster_label'):
    print(f"Cluster {lab}: {len(sub_df)} samples")
    for i, s in enumerate(sub_df['Sample']):
        print(f"  {s:19}", end="")
        if i % n_cols == n_cols - 1:
            print()
    print("\n")
    

t = ""
for lab, sub in df.groupby('cluster_label'):
    t += f"<h2>Cluster {lab}</h2>"
    t += " ".join(sub["Sample"])
    t += sub.describe().to_html()
#display_html(t, raw=True)


# --------------------------------------------------------
# Plotting and Visualization
# --------------------------------------------------------

# Create hexagons for visualization
def hexagon(_xx, _yy, radius):
    # Returns a NX x NY x 6 x 2 array of hexagon coordinates
    hex_xy = np.zeros((_xx.shape[0], _xx.shape[1], 7, 2))
    for i in range(7):
        theta = np.pi / 6 + 2 * np.pi * i / 6
        hex_xy[:, :, i, 0] = _xx + radius * np.cos(theta)
        hex_xy[:, :, i, 1] = _yy + radius * np.sin(theta)
    return hex_xy


hexagons = hexagon(xx, yy, 0.5)
hexagons.shape

# Cluster polygons
node_hexagons = hexagon(xx, yy, np.sqrt(3) / 3 + 0.001).reshape(-1, 7, 2)
cluster_boundaries = {}
for c in np.unique(node_labels):
    matches = node_labels.flatten() == c
    joined_polygon = shapely.union_all([shapely.Polygon(p) for p in node_hexagons[matches]])
    cluster_boundaries[c] = joined_polygon

# Define functions for plotting
def add_cluster_boundaries(fig: go.Figure, row_index, col_index):
    for k, boundary in cluster_boundaries.items():
        
        if boundary.geom_type == 'Polygon':
            bnds = [boundary]
        elif boundary.geom_type == 'MultiPolygon':
            bnds = list(boundary.geoms)
        else:
            raise NotImplementedError(boundary.geom_type)
        
        for poly in bnds:
            if poly.boundary.geom_type == 'LineString':
                crds = np.array(poly.boundary.coords)
            else:
                crds = np.array(poly.boundary.geoms[0].coords)

            fig.add_scatter(
                x=crds[:, 0],
                y=crds[:, 1],
                # fill='toself',
                mode="lines",
                showlegend=False,
                line=dict(color='black', width=2),
                row=row_index,
                col=col_index,
            )

def add_som_features(fig: go.Figure, row, col, feature_index, colorbar_settings):
    if feature_index >= weights.shape[2]:
        w = umatrix.copy()
    else:
        # The value of the weights will refer to the transformed data. If we want
        # the original scale, we will need to apply the inverse transform here
        w = weights[:, :, feature_index]

    w_min = w.min()
    w_max = w.max()
    w -= w.min()
    w /= w.max()
    colors = np.array(
        sample_colorscale(features_colorscale, w.flatten())
    ).reshape(final_nx, final_ny)


    # Add colored hexagons
    for i in range(final_nx):
        for j in range(final_ny):
            text = str(sorted(df.iloc[win_map[(i, j)]].index.values.tolist()))

            fig.add_scatter(
                x=hexagons[i, j, :, 0],
                y=hexagons[i, j, :, 1],
                fill='toself',
                mode='lines',
                fillcolor=colors[i, j],
                line=dict(color='black', width=0),
                text=text,
                showlegend=False,
                hoverinfo="text",
                hoveron="fills",
                row=row,
                col=col,
            )
  #_____________________________________________________________________________________________________________________         
   # Es Vertical colorbar
    fig.add_scatter(
        x=[0, 0],
        y=[0, 0],
        showlegend=False,
        marker=dict(color=[w_min, w_max], colorscale=features_colorscale, size=0.001, colorbar=colorbar_settings),
        row=row,
        col=col,
    )


def add_training_cells(fig: go.Figure, row, col):
    # Add colored hexagons
    label_colors = DEFAULT_PLOTLY_COLORS
    for i in range(final_nx):
        for j in range(final_ny):
            text = ""
            if len(win_map[(i, j)]) > 0:
                samples = [str(s) for s in sorted(df.iloc[win_map[(i, j)]].index.values.tolist())]
                if len(samples) <= 3:
                    text = f"{','.join(samples)}"
                elif len(samples) <= 6:
                    text = f"{','.join(samples[:3])}" + f"<br>{','.join(samples[3:])}"
                elif len(samples) <= 9:
                    text = f"{','.join(samples[:3])}" + f"<br>{','.join(samples[3:6])}" + f"<br>{','.join(samples[6:])}"
                else:
                    text = f"{','.join(samples[:3])}" + f"<br>{','.join(samples[3:-3])}" + f"<br>{','.join(samples[-3:])}"

            fig.add_scatter(
                x=hexagons[i, j, :, 0],
                y=hexagons[i, j, :, 1],
                fill='toself',
                mode='lines',
                fillcolor=label_colors[node_labels[i, j]],
                line=dict(width=0),
                text=text,
                showlegend=False,
                hoverinfo="text",
                hoveron="fills",
                row=row,
                col=col,
            )

            mid = (hexagons[i, j, 0, :] + hexagons[i, j, 3, :]) / 2
            fig.add_scatter(
                x=[mid[0]],
                y=[mid[1]],
                mode="text",
                text=text,
                textfont=dict(color="black", size=10),
                marker=dict(color='black'),
                showlegend=False,
                hoverinfo="none",
                row=row,
                col=col,
            )


n_plots = len(input_columns) + 2
subplot_rows = (n_plots + 1) // 2
subplot_columns = 2
features_fig = make_subplots(
    rows=subplot_rows,
    cols=subplot_columns,
    subplot_titles=input_columns + ["Umatrix"] + ["Clusters and winning cells"],
    horizontal_spacing= 0.05,
    vertical_spacing=0.05,

)
features_fig.layout.height = 50 + 250 * subplot_rows
#features_fig.layout.width = 1200


subplot_ix = list(
    itertools.product(range(1, subplot_rows + 1), range(1, subplot_columns + 1))
)[:n_plots]

# Feature plots
for i, (row_index, col_index) in enumerate(subplot_ix[:len(input_columns)+1]):
    # add_cluster_boundaries(features_fig, row_index, col_index)
    colorbar_settings = dict(
        thickness = 20,
        len = 0.5 * (1 / subplot_rows),
        y = 1 - (row_index - 0.5) / subplot_rows,
        x = 0.38 if col_index == 1 else 0.90
    )
    add_som_features(features_fig, row_index, col_index, i, colorbar_settings)

add_training_cells(features_fig, subplot_ix[-1][0], subplot_ix[-1][1])
add_cluster_boundaries(features_fig, subplot_ix[-1][0], subplot_ix[-1][1])


for i, (row_index, col_index) in enumerate(subplot_ix):
    sp = features_fig.get_subplot(row_index, col_index)
    sp.yaxis.scaleanchor = sp.yaxis.anchor
    sp.yaxis.showticklabels = False
    sp.xaxis.showticklabels = False

   

## -------------------------------------------------------
    # Draw isolines for specific SOM planes
    import matplotlib.pyplot as plt

    def add_isolines_to_som(fig, som_weights, coords, hexagons, feature_name, isoline_values, row, col):
        """
        Add contour isolines to a SOM plane.
        
        Parameters:
        - fig: plotly figure object
        - som_weights: weight matrix for the specific feature
        - coords: euclidean coordinates (xx, yy)
        - hexagons: hexagon coordinates for plotting
        - feature_name: name of the feature
        - isoline_values: list of values for which to draw isolines
        - row, col: subplot position
        """
        xx, yy = coords
        
        # Create interpolated grid for smoother contours
        points = np.column_stack([xx.flatten(), yy.flatten()])
        values = som_weights.flatten()
        
        # Create fine grid
        grid_x = np.linspace(xx.min(), xx.max(), 100)
        grid_y = np.linspace(yy.min(), yy.max(), 100)
        grid_xx, grid_yy = np.meshgrid(grid_x, grid_y)
        
        # Interpolate
        grid_z = griddata(points, values, (grid_xx, grid_yy), method='cubic')
        
        # Generate contours using matplotlib (then extract and plot in plotly)
        fig_temp = plt.figure()
        ax_temp = fig_temp.add_subplot(111)
        contours = ax_temp.contour(grid_xx, grid_yy, grid_z, levels=isoline_values, colors='black', linewidths=2)
        plt.close(fig_temp)
      
        
        # Extract contour paths and add to plotly figure
        for level_idx in range(len(contours.levels)):
            paths = contours.get_paths()[level_idx] if hasattr(contours, 'get_paths') else contours.allsegs[level_idx]
            if not isinstance(paths, list):
                paths = [paths]
            for path in paths:
                vertices = path.vertices if hasattr(path, 'vertices') else path
                fig.add_scatter(
                    x=vertices[:, 0],
                    y=vertices[:, 1],
                    mode='lines',
                    line=dict(color='darkgreen', width=3.0, dash='dash'),
                    showlegend=False,
                    hoverinfo='skip',
                    row=row,
                    col=col
                )

    
   
    # Define isoline values for each feature
    isoline_config = {
              
        'F4': [0.2, 0.25, 0.36, 0.4, 0.45, 0.5, 0.55, 0.6, 0.7, 0.8, 0.85, 0.9, 1.0],
        'F5': [0.1, 0.2, 0.3,  0.5, 0.55, 0.6, 0.65, 0.7, 0.75, 0.8],
        'F7': [0.1,  0.2, 0.25, 0.3, 0.32, 0.4, 0.45, 0.5, 0.6, 0.7, 0.8, 0.9],
        'F8': [0.1, 0.2, 0.3, 0.35, 0.44, 0.5, 0.55, 0.6, 0.65, 0.7],
        'F9': [0.1, 0.2, 0.25, 0.3, 0.4, 0.55, 0.6, 0.65, 0.7, 0.72, 0.75],
       
    }
 
    # Add isolines to the features_fig (main SOM visualization)
    if 'features_fig' in locals() and 'weights' in locals():
        for feat_name, isoline_vals in isoline_config.items():
            if feat_name in input_columns:
                feat_idx = input_columns.index(feat_name)
                # Find subplot position
                plot_position = feat_idx
                row_idx = subplot_ix[plot_position][0]
                col_idx = subplot_ix[plot_position][1]
                
                add_isolines_to_som(
                    features_fig,
                    weights[:, :, feat_idx],
                    (xx, yy),
                    hexagons,
                    feat_name,
                    isoline_vals,
                    row_idx,
                    col_idx
                )

# Define isoline colors for specific values
    isoline_colors = {
        #'F3': {0.38: 'darkorange'},
        'F4': {0.31: 'darkorange'},
        'F5': {0.59: 'darkorange', 0.72: 'darkorange'},
        'F7': {0.14: 'darkorange'},
        'F8': {0.37: 'darkorange'},
        'F9': {0.48: 'darkorange'},  
    }

    # Define isoline styles for specific values
    isoline_styles = {
        #'F3': {0.38: 'solid'},
        'F4': {0.31: 'solid'},
        'F5': {0.59: 'solid', 0.72: 'dashdot'}, 
        'F7': {0.14: 'solid'},
        'F8': {0.37: 'solid'},
        'F9': {0.48: 'solid'},
    }



    # Add colored isolines to the features_fig (main SOM visualization)
    if 'features_fig' in locals() and 'weights' in locals():
        for feat_name, color_dict in isoline_colors.items():
            if feat_name in input_columns:
                feat_idx = input_columns.index(feat_name)
                # Find subplot position
                plot_position = feat_idx
                row_idx = subplot_ix[plot_position][0]
                col_idx = subplot_ix[plot_position][1]
                
                # Get the weight matrix for this feature
                som_weights = weights[:, :, feat_idx]
                
                # Create interpolated grid for smoother contours
                points = np.column_stack([xx.flatten(), yy.flatten()])
                values = som_weights.flatten()
                
                # Create fine grid
                grid_x = np.linspace(xx.min(), xx.max(), 100)
                grid_y = np.linspace(yy.min(), yy.max(), 100)
                grid_xx, grid_yy = np.meshgrid(grid_x, grid_y)
                
                # Interpolate
                grid_z = griddata(points, values, (grid_xx, grid_yy), method='cubic')
                
                # Draw isolines for specific values with specified colors
                for isoline_val, color in color_dict.items():
                    fig_temp = plt.figure()
                    ax_temp = fig_temp.add_subplot(111)
                    contours = ax_temp.contour(grid_xx, grid_yy, grid_z, levels=[isoline_val], colors='black', linewidths=3.0)
                    plt.close(fig_temp)
                    
                    # Get the line style for this isoline value
                    line_style = 'solid'  # default
                    if feat_name in isoline_styles and isoline_val in isoline_styles[feat_name]:
                        line_style = isoline_styles[feat_name][isoline_val]
                    
                    # Extract contour paths and add to plotly figure
                    for level_idx in range(len(contours.levels)):
                        paths = contours.allsegs[level_idx]
                        for path in paths:
                            features_fig.add_scatter(
                                x=path[:, 0],
                                y=path[:, 1],
                                mode='lines',
                                line=dict(color=color, width=4.0, dash=line_style),
                                showlegend=False,
                                hoverinfo='skip',
                                row=row_idx,
                                col=col_idx
                            )
    # Add isolines to the features_fig (main SOM visualization)
    if 'features_fig' in locals() and 'weights' in locals():
        for feat_name, isoline_vals in isoline_config.items():
            if feat_name in input_columns:
                feat_idx = input_columns.index(feat_name)
                # Find subplot position
                plot_position = feat_idx
                row_idx = subplot_ix[plot_position][0]
                col_idx = subplot_ix[plot_position][1]
                
                add_isolines_to_som(
                    features_fig,
                    weights[:, :, feat_idx],
                    (xx, yy),
                    hexagons,
                    feat_name,
                    isoline_vals,
                    row_idx,
                    col_idx
                )

                # Calculate PC1 value for each neuron (same across all features)
                # PC1_neuron = sum(wji * Li) where wji are neuron weights, Li are PC1 loadings
                
                pc1_neuron_map = np.zeros((final_nx, final_ny))
                
                for i in range(final_nx):
                    for j in range(final_ny):
                        # Get weights for this neuron across all features
                        neuron_weights = weights[i, j, :]
                        
                        # Calculate PC1 value as dot product of neuron weights and PC1 loadings
                        pc1_neuron_map[i, j] = np.dot(neuron_weights, pca.components_[0, :])
                
                # Use mean PC1 value as the isoline threshold
                pc1_isoline_value = np.mean(pc1_neuron_map)

                print(f"\nPC1 Isoline Value (mean): {pc1_isoline_value:.6f}")
                print(f"PC1 range: [{pc1_neuron_map.min():.6f}, {pc1_neuron_map.max():.6f}]")
                print("="*50)

                # Add PC1 isolines to each feature plot (using PC1 neuron map)
                def add_pc1_isolines(fig, coords, hexagons, feature_name, row, col):
                    """
                    Add PC1 isoline to a SOM plane based on PC1 neuron values.
                    This ensures the same geometric line across all feature plots.
                    """
                    xx, yy = coords
                    
                    # Create interpolated grid for smoother contours using PC1 neuron map
                    points = np.column_stack([xx.flatten(), yy.flatten()])
                    values = pc1_neuron_map.flatten()
                    
                    # Create fine grid
                    grid_x = np.linspace(xx.min(), xx.max(), 100)
                    grid_y = np.linspace(yy.min(), yy.max(), 100)
                    grid_xx, grid_yy = np.meshgrid(grid_x, grid_y)
                    
                    # Interpolate PC1 values
                    grid_z = griddata(points, values, (grid_xx, grid_yy), method='cubic')
                    
                    # Generate contour for PC1 isoline value
                    fig_temp = plt.figure()
                    ax_temp = fig_temp.add_subplot(111)
                    contours = ax_temp.contour(grid_xx, grid_yy, grid_z, levels=[pc1_isoline_value], colors='black', linewidths=2)
                    plt.close(fig_temp)
                    
                    # Extract contour paths and add to plotly figure
                    if len(contours.allsegs) > 0:
                        for path in contours.allsegs[0]:
                            fig.add_scatter(
                                x=path[:, 0],
                                y=path[:, 1],
                                mode='lines',
                                line=dict(color='red', width=4.5, dash='dot'),
                                showlegend=False,
                                hoverinfo='skip',
                                row=row,
                                col=col
                            )

                # Add PC1 isolines to all feature plots
                pc1_features = input_columns.copy()
                # if 'pH_L' in pc1_features:
                #     pc1_features.remove('pH_L')  # Exclude pH_L if present

                for feat_name in pc1_features:
                    if feat_name in input_columns:
                        feat_idx = input_columns.index(feat_name)
                        plot_position = feat_idx
                        row_idx = subplot_ix[plot_position][0]
                        col_idx = subplot_ix[plot_position][1]
                        
                        add_pc1_isolines(
                            features_fig,
                            (xx, yy),
                            hexagons,
                            feat_name,
                            row_idx,
                            col_idx
                        )

                # print(f"PC1 isolines added to feature plots (excluding pH_L)")
                # print("="*50 + "\n")

                # Print PC1 value for each neuron
                print("\n" + "="*50)
                print("PC1 Value for Each Neuron")
                print("="*50)

                # Calculate PC1 value for each neuron
                pc1_neuron_values = np.zeros((final_nx, final_ny))

                for i in range(final_nx):
                    for j in range(final_ny):
                        # Get weights for this neuron across all features
                        neuron_weights = weights[i, j, :]
                        
                        # Calculate PC1 value as dot product of neuron weights and PC1 loadings
                        pc1_value = np.dot(neuron_weights, pca.components_[0, :])
                        pc1_neuron_values[i, j] = pc1_value
                        
                        #print(f"Neuron ({i}, {j}): PC1 = {pc1_value:.6f}")

                print(f"\nPC1 neuron values - Min: {pc1_neuron_values.min():.6f}, Max: {pc1_neuron_values.max():.6f}, Mean: {pc1_neuron_values.mean():.6f}")
                print("="*50 + "\n")

  


# Calculate PC1 general isoline value (it must be the same for all features) by equation: (sum(sum(wji*Li))/K), where K is the amount of neurons, wji are the weights of neurons, Li are Lodings of PCA. Construct PC1 isoline on the each feature plot as red dotted line.

 
# save dataframe as CSV
# df.to_csv('myoutputSil.csv')

# df = pd.DataFrame(df)
# df['log_cluster_label'] = np.log(df['cluster_label'])
# df.to_csv('myoutputforTDS.csv')


# Save dataframe as CSV in specified folder
output_folder = Path('H:/Python/Self_Org_Map/As_Progr_for-Som/A0_Es_Final_2')
output_folder.mkdir(parents=True, exist_ok=True)

df = pd.DataFrame(df)
df['log_cluster_label'] = np.log(df['cluster_label'])
output_file = output_folder / 'myoutputforTDS.csv'
df.to_csv(output_file, index=False)



# Save each cluster to a separate CSV file
output_dir = Path(__file__).parent / "cluster_after_factors_forTDS"
output_dir.mkdir(parents=True, exist_ok=True)

unique_labels = sorted(df['cluster_label'].dropna().unique())
for cl in unique_labels:
    sub_df = df[df['cluster_label'] == cl]
    out_path = output_dir / f"myoutputSil_cluster_label_{int(cl)}.csv"
    sub_df.to_csv(out_path, index=False)

    

 # Increase font size of numbers on horizontal colorbars
    for _fig in [locals().get('features_fig')]:
        if _fig:
            for tr in _fig.data:
                if hasattr(tr, 'marker') and getattr(tr.marker, 'colorbar', None):
                    tr.marker.colorbar.update(tickfont=dict(size=18))

#Shows SOM planes with vertical scale bar
show_figure(_fig)




 # Force all horizontal colorbars (fig_new / fig_sep) to span 0..1 with plt.stepstep 0.0, 0.2, 0.4, 0.6, 0.8, 1.0
for _fig in [locals().get('features_fig')]:
    if not _fig:
        continue
    for tr in _fig.data:
        mk = getattr(tr, 'marker', None)
        cb = getattr(mk, 'colorbar', None)
        if mk and cb:
            cb.update(
                orientation='h',
                thickness=15,
                len=0.25, 
                x=0.5,
                xanchor='center',
                y=-0.15,
                yanchor='top',
                tickvals=[0.0, 0.2, 0.4, 0.6, 0.8, 1.0],
                ticktext=["0.0", "0.2", "0.4", "0.6", "0.8", "1.0"],
                ticks="outside",
                tickfont=dict(size=18)
            )

#Shows SOM planes with horizontal scale bar
#show_figure(features_fig)


# Calculate and print Ward's distance for each feature
print("\n" + "="*50)
print("Ward's Distance for Each Feature")
print("="*50)

for feature_name in ['F3', 'F4', 'F5', 'F7', 'F8', 'F9']:
    if feature_name in input_columns:
        feat_idx = input_columns.index(feature_name)
        feature_weights = weights[:, :, feat_idx].reshape(-1, 1)
        
        # Calculate Ward's linkage for this feature
        ward_linkage = hierarchy.linkage(feature_weights, method='ward', metric='euclidean')
        
        # The maximum distance in the linkage matrix represents Ward's distance
        max_ward_distance = ward_linkage[:, 2].max()
        
        print(f"{feature_name}: {max_ward_distance:.4f}")

print("="*50 + "\n")

# Calculate and print variance values for each feature
print("\n" + "="*50)
print("Variance Values for Each Feature")
print("="*50)

for feature_name in ['F3', 'F4', 'F5', 'F7', 'F8', 'F9']:
    if feature_name in input_columns:
        feat_idx = input_columns.index(feature_name)
        # Get the feature data from the transformed data
        #feature_data = transformed_data[:, feat_idx]
        feature_data = weights[:, :, feat_idx].reshape(-1)
        
        # Calculate variance
        #variance = np.var(feature_data)
        variance = np.var(feature_data)
        
        print(f"{feature_name}: {variance:.6f}")

print("="*50 + "\n")



# Calculate and print weight ranges for each feature
print("\n" + "="*50)
print("Weight Ranges for Each Feature")
print("="*50)

for feature_name in ['F3', 'F4', 'F5', 'F7', 'F8', 'F9']:
    if feature_name in input_columns:
        feat_idx = input_columns.index(feature_name)
        # Get the feature weights from the SOM
        feature_weights = weights[:, :, feat_idx].flatten()
        
        # Calculate min, max, and mean
        mean_weight = feature_weights.mean()
        
        print(f"{feature_name}: Mean={mean_weight:.4f}")

print("="*50 + "\n")




# Calculate and print correlation matrix between features
print("\n" + "="*50)
print("Correlation Matrix Between Features")
print("="*50)

# Define the features to analyze
correlation_features = ['F3', 'F4', 'F5', 'F7', 'F8', 'F9']

# Check which features are available in input_columns
available_features = [f for f in correlation_features if f in input_columns]

if len(available_features) > 1:
    # Extract the feature indices
    feature_indices = [input_columns.index(f) for f in available_features]
    
    # Get the data for these features from transformed_data
    feature_data = transformed_data[:, feature_indices]
    
    # Calculate correlation matrix
    correlation_matrix = np.corrcoef(feature_data.T)
    
    # Create a DataFrame for better visualization
    corr_df = pd.DataFrame(
        correlation_matrix,
        index=available_features,
        columns=available_features
    )
    
    print("\nCorrelation Matrix:")
    print(corr_df.to_string())
    
    # Create a heatmap using plotly
    fig_corr = go.Figure(data=go.Heatmap(
        z=correlation_matrix,
        x=available_features,
        y=available_features,
        colorscale='RdBu',
        zmid=0,
        text=correlation_matrix,
        texttemplate='%{text:.3f}',
        textfont={"size": 12},
        colorbar=dict(title="Correlation")
    ))
    
    fig_corr.update_layout(
        title='Feature Correlation Matrix',
        xaxis_title='Features',
        yaxis_title='Features',
        width=700,
        height=600
    )
    
    #show_figure(fig_corr)
    
    # Save correlation matrix to CSV
    corr_df.to_csv('feature_correlation_matrix.csv')
    print("\nCorrelation matrix saved to 'feature_correlation_matrix.csv'")
else:
    print("Not enough features available for correlation analysis")

print("="*50 + "\n")



# # Save weight values for each neuron of specified features to separate Excel files
# print("\n" + "="*50)
# print("Saving Neuron Weight Values to Excel Files")
# print("="*50)

features_to_save = ['F3', 'F4', 'F5', 'F7', 'F8', 'F9']

for feature_name in features_to_save:
    if feature_name in input_columns:
        feat_idx = input_columns.index(feature_name)
        
        # Get the weight matrix for this feature
        feature_weights = weights[:, :, feat_idx]
        
        # Create DataFrame with neuron coordinates
        neuron_data = []
        for i in range(final_nx):
            for j in range(final_ny):
                neuron_data.append({
                    'Neuron_X': i,
                    'Neuron_Y': j,
                    'Weight': feature_weights[i, j]
                })
        
        neuron_df = pd.DataFrame(neuron_data)
        
        # Save to Excel file
        output_filename = f'neuron_weights_{feature_name}.xlsx'
        neuron_df.to_excel(output_filename, index=False)
        
#         print(f"Saved {feature_name} neuron weights to '{output_filename}'")

# print("="*50 + "\n")

# Save weight values for each neuron of specified features to separate Excel files
print("\n" + "="*50)
print("Saving Neuron Weight Values to Excel Files")
print("="*50)

features_to_save = ['F3', 'F4', 'F5', 'F7', 'F8', 'F9']

# Create output directory for neuron weights
neuron_weights_dir = Path('H:/Python/Self_Org_Map/As_Progr_for-Som/A0_Es_Final_2/neuron_weights')
neuron_weights_dir.mkdir(parents=True, exist_ok=True)

for feature_name in features_to_save:
    if feature_name in input_columns:
        feat_idx = input_columns.index(feature_name)
        
        # Get the weight matrix for this feature
        feature_weights = weights[:, :, feat_idx]
        
        # Create DataFrame with neuron coordinates
        neuron_data = []
        for i in range(final_nx):
            for j in range(final_ny):
                neuron_data.append({
                    'Neuron_X': i,
                    'Neuron_Y': j,
                    'Weight': feature_weights[i, j]
                })
        
        neuron_df = pd.DataFrame(neuron_data)
        
        # Save to Excel file in the specified folder
        output_filename = neuron_weights_dir / f'neuron_weights_{feature_name}.xlsx'
        neuron_df.to_excel(output_filename, index=False)
        
#         print(f"Saved {feature_name} neuron weights to '{output_filename}'")

# print("="*50 + "\n")

# # Create biplot with loading vectors
# print("\n" + "="*50)
# print("PCA Biplot: Samples and Feature Loading Vectors")
# print("="*50)

# Perform PCA on the transformed data
pca_full = PCA(n_components=2)
pca_transformed = pca_full.fit_transform(transformed_data)

# Get the loadings (principal components)
loadings = pca_full.components_.T * np.sqrt(pca_full.explained_variance_)

# Create the biplot
fig_biplot = go.Figure()

# Add scatter plot for samples
fig_biplot.add_scatter(
    x=pca_transformed[:, 0],
    y=pca_transformed[:, 1],
    mode='markers',
    marker=dict(
        size=10,
        color=df['cluster_label'],
        colorscale='Viridis',
        showscale=True,
        colorbar=dict(title="Cluster", x=1.15)
    ),
    text=df['Sample'],
    name='Samples',
    hovertemplate='<b>%{text}</b><br>PC1: %{x:.3f}<br>PC2: %{y:.3f}<extra></extra>'
)

# Add loading vectors
scale_factor = 3.5  # Adjust this to make vectors visible
for i, feature in enumerate(input_columns):
    fig_biplot.add_annotation(
        ax=0.0, ay=0.0,
        axref='x', ayref='y',
        x=loadings[i, 0] * scale_factor,
        y=loadings[i, 1] * scale_factor,
        xref='x', yref='y',
        showarrow=True,
        arrowhead=2,
        arrowsize=1,
        arrowwidth=2,
        arrowcolor='red',
    )
    
    # Add feature labels at the end of vectors
    fig_biplot.add_annotation(
        x=loadings[i, 0] * scale_factor * 1.1,
        y=loadings[i, 1] * scale_factor * 1.1,
        xref='x', yref='y',
        text=feature,
        showarrow=False,
        font=dict(size=12, color='red', family='Arial Black'),
        bgcolor='rgba(255, 255, 255, 0.8)',
        #bordercolor='red',
        #borderwidth=1
    )

fig_biplot.update_layout(
    title=f'PCA Biplot - Samples and Feature Loadings',
    xaxis_title=f'PC1 ({pca_full.explained_variance_ratio_[0]*100:.1f}%)',
    yaxis_title=f'PC2 ({pca_full.explained_variance_ratio_[1]*100:.1f}%)',
    width=900,
    height=700,
    showlegend=True
)

#show_figure(fig_biplot)

# # Print loading values
# print("\nFeature Loadings on PC1 and PC2:")
# print("-" * 40)
# for i, feature in enumerate(input_columns):
#     print(f"{feature:10s}: PC1={loadings[i, 0]:7.4f}, PC2={loadings[i, 1]:7.4f}")

# print("="*50 + "\n")

# # Save weight values for each neuron of specified features to separate Excel files
# print("\n" + "="*50)
# print("Saving Neuron Weight Values to Excel Files")
# print("="*50)

features_to_save = ['F3', 'F4', 'F5', 'F7', 'F8', 'F9']

for feature_name in features_to_save:
    if feature_name in input_columns:
        feat_idx = input_columns.index(feature_name)
        
        # Get the weight matrix for this feature
        feature_weights = weights[:, :, feat_idx]
        
        # Create DataFrame with neuron coordinates
        neuron_data = []
        for i in range(final_nx):
            for j in range(final_ny):
                neuron_data.append({
                    'Neuron_X': i,
                    'Neuron_Y': j,
                    'Weight': feature_weights[i, j]
                })
        
        neuron_df = pd.DataFrame(neuron_data)
        
        # Save to Excel file
        output_filename = f'neuron_weights_{feature_name}.xlsx'
        neuron_df.to_excel(output_filename, index=False)
        
#         print(f"Saved {feature_name} neuron weights to '{output_filename}'")

# print("="*50 + "\n")

# # Plot boundaries between clusters on each SOM plane
# print("\n" + "="*50)
# print("Creating SOM Planes with Cluster Boundaries")
# print("="*50)

# Create a new figure with subplots for each feature + U-matrix + clusters
n_plots_with_boundaries = len(input_columns) + 2
subplot_rows_boundaries = (n_plots_with_boundaries + 1) // 2
subplot_columns_boundaries = 2

fig_with_boundaries = make_subplots(
    rows=subplot_rows_boundaries,
    cols=subplot_columns_boundaries,
    subplot_titles=input_columns + ["U-Matrix"] + ["Clusters"],
    horizontal_spacing=0.05,
    vertical_spacing=0.05,
)
fig_with_boundaries.layout.height = 50 + 250 * subplot_rows_boundaries

subplot_ix_boundaries = list(
    itertools.product(range(1, subplot_rows_boundaries + 1), range(1, subplot_columns_boundaries + 1))
)[:n_plots_with_boundaries]

# Feature plots with cluster boundaries
for i, (row_index, col_index) in enumerate(subplot_ix_boundaries[:len(input_columns)+1]):
    colorbar_settings = dict(
        thickness=20,
        len=0.5 * (1 / subplot_rows_boundaries),
        y=1 - (row_index - 0.5) / subplot_rows_boundaries,
        x=0.38 if col_index == 1 else 0.90
    )
    add_som_features(fig_with_boundaries, row_index, col_index, i, colorbar_settings)
    add_cluster_boundaries(fig_with_boundaries, row_index, col_index)

# Add clusters plot
add_training_cells(fig_with_boundaries, subplot_ix_boundaries[-1][0], subplot_ix_boundaries[-1][1])
add_cluster_boundaries(fig_with_boundaries, subplot_ix_boundaries[-1][0], subplot_ix_boundaries[-1][1])

# Update axes
for i, (row_index, col_index) in enumerate(subplot_ix_boundaries):
    sp = fig_with_boundaries.get_subplot(row_index, col_index)
    sp.yaxis.scaleanchor = sp.yaxis.anchor
    sp.yaxis.showticklabels = False
    sp.xaxis.showticklabels = False

# Update colorbar font size
for tr in fig_with_boundaries.data:
    if hasattr(tr, 'marker') and getattr(tr.marker, 'colorbar', None):
        tr.marker.colorbar.update(tickfont=dict(size=18))

show_figure(fig_with_boundaries)

# print("SOM planes with cluster boundaries displayed")
# print("="*50 + "\n")


# # Plot boundaries between clusters on each SOM plane with isolines
# print("\n" + "="*50)
# print("Creating SOM Planes with Cluster Boundaries and Isolines")
# print("="*50)

# Create a new figure with subplots for each feature + U-matrix + clusters
n_plots_with_boundaries_isolines = len(input_columns) + 2
subplot_rows_boundaries_isolines = (n_plots_with_boundaries_isolines + 1) // 2
subplot_columns_boundaries_isolines = 2

fig_with_boundaries_isolines = make_subplots(
    rows=subplot_rows_boundaries_isolines,
    cols=subplot_columns_boundaries_isolines,
    subplot_titles=input_columns + ["U-Matrix"] + ["Clusters"],
    horizontal_spacing=0.05,
    vertical_spacing=0.05,
)
fig_with_boundaries_isolines.layout.height = 50 + 250 * subplot_rows_boundaries_isolines

subplot_ix_boundaries_isolines = list(
    itertools.product(range(1, subplot_rows_boundaries_isolines + 1), range(1, subplot_columns_boundaries_isolines + 1))
)[:n_plots_with_boundaries_isolines]

# Feature plots with cluster boundaries and isolines
for i, (row_index, col_index) in enumerate(subplot_ix_boundaries_isolines[:len(input_columns)+1]):
    colorbar_settings = dict(
        thickness=20,
        len=0.5 * (1 / subplot_rows_boundaries_isolines),
        y=1 - (row_index - 0.5) / subplot_rows_boundaries_isolines,
        x=0.38 if col_index == 1 else 0.90
    )
    add_som_features(fig_with_boundaries_isolines, row_index, col_index, i, colorbar_settings)
    add_cluster_boundaries(fig_with_boundaries_isolines, row_index, col_index)
    
    # Add isolines for specific features
    if i < len(input_columns):
        feat_name = input_columns[i]
        if feat_name in isoline_config:
            add_isolines_to_som(
                fig_with_boundaries_isolines,
                weights[:, :, i],
                (xx, yy),
                hexagons,
                feat_name,
                isoline_config[feat_name],
                row_index,
                col_index
            )
        
        # Add PC1 isolines
        if feat_name in pc1_features:
            add_pc1_isolines(
                fig_with_boundaries_isolines,
                (xx, yy),
                hexagons,
                feat_name,
                row_index,
                col_index
            )

# Add clusters plot
add_training_cells(fig_with_boundaries_isolines, subplot_ix_boundaries_isolines[-1][0], subplot_ix_boundaries_isolines[-1][1])
add_cluster_boundaries(fig_with_boundaries_isolines, subplot_ix_boundaries_isolines[-1][0], subplot_ix_boundaries_isolines[-1][1])

# Update axes
for i, (row_index, col_index) in enumerate(subplot_ix_boundaries_isolines):
    sp = fig_with_boundaries_isolines.get_subplot(row_index, col_index)
    sp.yaxis.scaleanchor = sp.yaxis.anchor
    sp.yaxis.showticklabels = False
    sp.xaxis.showticklabels = False

# Update colorbar font size
for tr in fig_with_boundaries_isolines.data:
    if hasattr(tr, 'marker') and getattr(tr.marker, 'colorbar', None):
        tr.marker.colorbar.update(tickfont=dict(size=18))

show_figure(fig_with_boundaries_isolines)

# print("SOM planes with cluster boundaries and isolines displayed")
# print("="*50 + "\n")


#Write on the separate word dokument (save in the folder H:\Python\Self_Org_Map\As_Progr_for-Som\A0_Es_Final_2) intepretation of SOM featrures and clusters based on the above analysis. Consider also the PCA biplot and correlation matrix results. Consider cases of movement of the siolines with darkorange color and PC1 isolines. Discuss what the clusters might represent in terms of the original features.

# # Create Word document with SOM interpretation
# print("\n" + "="*50)
# print("Creating Word Document with SOM Interpretation")
# print("="*50)

doc = Document()

# Add title
title = doc.add_heading('SOM Feature and Cluster Interpretation', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Add metadata section
doc.add_heading('Analysis Metadata', level=1)
p = doc.add_paragraph()
p.add_run(f'Date of Analysis: {pd.Timestamp.now().strftime("%Y-%m-%d %H:%M:%S")}\n').bold = True
p.add_run(f'Number of Samples: {len(df)}\n')
p.add_run(f'Number of Features: {len(input_columns)}\n')
p.add_run(f'Features Analyzed: {", ".join(input_columns)}\n')
p.add_run(f'SOM Dimensions: {final_nx} x {final_ny}\n')
p.add_run(f'Number of Clusters: {int(selected_cluster["Number of clusters"])}\n')
p.add_run(f'Quantization Error: {som.quantization_error(transformed_data):.4f}\n')
p.add_run(f'Topographic Error: {som.topographic_error(transformed_data):.4f}\n')

# PCA Summary
doc.add_heading('PCA Summary', level=1)
p = doc.add_paragraph()
p.add_run(f'PC1 Explained Variance: {pca_full.explained_variance_ratio_[0]*100:.2f}%\n')
p.add_run(f'PC2 Explained Variance: {pca_full.explained_variance_ratio_[1]*100:.2f}%\n')
p.add_run(f'Total Variance Explained: {(pca_full.explained_variance_ratio_[0] + pca_full.explained_variance_ratio_[1])*100:.2f}%\n')

# Feature Loadings
doc.add_heading('Feature Loadings on Principal Components', level=2)
table = doc.add_table(rows=1, cols=3)
table.style = 'Light Grid Accent 1'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Feature'
hdr_cells[1].text = 'PC1 Loading'
hdr_cells[2].text = 'PC2 Loading'

for i, feature in enumerate(input_columns):
    row_cells = table.add_row().cells
    row_cells[0].text = feature
    row_cells[1].text = f'{loadings[i, 0]:.4f}'
    row_cells[2].text = f'{loadings[i, 1]:.4f}'

# Feature Statistics
doc.add_heading('Feature Statistics (SOM Weights)', level=1)
table = doc.add_table(rows=1, cols=4)
table.style = 'Light Grid Accent 1'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Feature'
hdr_cells[1].text = 'Mean Weight'
hdr_cells[2].text = 'Variance'
hdr_cells[3].text = 'Ward Distance'

for feature_name in input_columns:
    if feature_name in ['F3', 'F4', 'F5', 'F7', 'F8', 'F9']:
        feat_idx = input_columns.index(feature_name)
        feature_weights = weights[:, :, feat_idx].flatten()
        feature_data = weights[:, :, feat_idx].reshape(-1)
        feature_weights_2d = weights[:, :, feat_idx].reshape(-1, 1)
        
        mean_weight = feature_weights.mean()
        variance = np.var(feature_data)
        ward_linkage = hierarchy.linkage(feature_weights_2d, method='ward', metric='euclidean')
        max_ward_distance = ward_linkage[:, 2].max()
        
        row_cells = table.add_row().cells
        row_cells[0].text = feature_name
        row_cells[1].text = f'{mean_weight:.4f}'
        row_cells[2].text = f'{variance:.6f}'
        row_cells[3].text = f'{max_ward_distance:.4f}'

# Cluster Distribution
doc.add_heading('Cluster Distribution', level=1)
table = doc.add_table(rows=1, cols=2)
table.style = 'Light Grid Accent 1'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Cluster'
hdr_cells[1].text = 'Number of Samples'

for lab in sorted(df['cluster_label'].dropna().unique()):
    count = len(df[df['cluster_label'] == lab])
    row_cells = table.add_row().cells
    row_cells[0].text = f'Cluster {int(lab)}'
    row_cells[1].text = str(count)

# Interpretation
doc.add_heading('Interpretation of SOM Features and Clusters', level=1)

doc.add_heading('1. PCA Biplot Analysis', level=2)
p = doc.add_paragraph()
p.add_run('The PCA biplot reveals the relationships between features and their contribution to sample variance:\n\n')

# Analyze PC1 and PC2 dominant features
pc1_loadings = [(input_columns[i], loadings[i, 0]) for i in range(len(input_columns))]
pc2_loadings = [(input_columns[i], loadings[i, 1]) for i in range(len(input_columns))]
pc1_loadings.sort(key=lambda x: abs(x[1]), reverse=True)
pc2_loadings.sort(key=lambda x: abs(x[1]), reverse=True)

p.add_run(f'• PC1 (dominant features): {", ".join([f[0] for f in pc1_loadings[:3]])}\n')
p.add_run(f'• PC2 (dominant features): {", ".join([f[0] for f in pc2_loadings[:3]])}\n')
p.add_run(f'\nFeatures with high positive PC1 loadings indicate variables that increase together along the first principal component. ')
p.add_run(f'Features with high PC2 loadings represent the secondary pattern of variation orthogonal to PC1.\n')

doc.add_heading('2. Correlation Analysis', level=2)
if len(available_features) > 1:
    p = doc.add_paragraph()
    p.add_run('Key correlations between features:\n\n')
    
    # Find strong correlations
    correlation_matrix = np.corrcoef(transformed_data[:, [input_columns.index(f) for f in available_features]].T)
    for i in range(len(available_features)):
        for j in range(i+1, len(available_features)):
            corr_val = correlation_matrix[i, j]
            if abs(corr_val) > 0.5:
                p.add_run(f'• {available_features[i]} and {available_features[j]}: {corr_val:.3f}')
                if corr_val > 0:
                    p.add_run(' (positive correlation)\n')
                else:
                    p.add_run(' (negative correlation)\n')

doc.add_heading('3. Isoline Analysis', level=2)
p = doc.add_paragraph()
p.add_run('Darkorange Isolines:\n')
p.add_run('These critical isolines represent threshold values that separate distinct hydrochemical regimes:\n\n')

if 'isoline_colors' in locals():
    for feat_name, color_dict in isoline_colors.items():
        if feat_name in input_columns:
            feat_idx = input_columns.index(feat_name)
            for isoline_val in color_dict.keys():
                p.add_run(f'• {feat_name} = {isoline_val:.2f}: ')
                p.add_run('Marks a significant boundary in the SOM space where this feature undergoes notable change.\n')

p.add_run('\nPC1 Isolines (red dotted):\n')
p.add_run('These isolines indicate where the first principal component contributes significantly, ')
p.add_run('representing transition zones between different dominant patterns in the data.\n')

doc.add_heading('4. Cluster Interpretation', level=2)
p = doc.add_paragraph()
p.add_run('Based on the analysis, the clusters likely represent distinct hydrochemical facies:\n\n')

for lab in sorted(df['cluster_label'].dropna().unique()):
    cluster_df = df[df['cluster_label'] == lab]
    p.add_run(f'Cluster {int(lab)} ({len(cluster_df)} samples):\n').bold = True
    
    # Calculate mean values for each feature in this cluster
    cluster_means = {}
    for col in input_columns:
        if col in cluster_df.columns:
            cluster_means[col] = cluster_df[col].mean()
    
    # Determine characteristic features
    p.add_run(f'  Characterized by: ')
    high_features = [f for f, v in cluster_means.items() if v > np.percentile(df[f], 66)]
    low_features = [f for f, v in cluster_means.items() if v < np.percentile(df[f], 33)]
    
    if high_features:
        p.add_run(f'High {", ".join(high_features[:3])}')
    if low_features:
        if high_features:
            p.add_run('; ')
        p.add_run(f'Low {", ".join(low_features[:3])}')
    p.add_run('\n')
    
    # List samples
    samples_list = cluster_df['Sample'].tolist()[:10]
    p.add_run(f'  Example samples: {", ".join(map(str, samples_list))}')
    if len(cluster_df) > 10:
        p.add_run(f' ... and {len(cluster_df)-10} more')
    p.add_run('\n\n')



# # Print sum of weights and PC1 loadings for each feature
# print("\n" + "="*50)
# print("Sum of Weights and PC1 Loadings for Each Feature")
# print("="*50)

for feature_name in input_columns:
    feat_idx = input_columns.index(feature_name)
    
    # Get the weight matrix for this feature
    feature_weights = weights[:, :, feat_idx]
    
    # Calculate sum of weights
    sum_weights = np.sum(feature_weights)
    
    # Get PC1 loading for this feature
    pc1_loading = pca.components_[0, feat_idx]
    
#    print(f"{feature_name:10s}: Sum of Weights={sum_weights:8.4f}, PC1 Loading={pc1_loading:7.4f}")

# print("="*50 + "\n")



# Create enhanced biplot with improved visualization
print("\n" + "="*50)
print("Creating Enhanced PCA Biplot")
print("="*50)

# Perform PCA on the transformed data (already done above, but ensure it's available)
pca_biplot = PCA(n_components=2)
pca_scores = pca_biplot.fit_transform(transformed_data)

# Get the loadings (principal components)
loadings_biplot = pca_biplot.components_.T * np.sqrt(pca_biplot.explained_variance_)

# Create the enhanced biplot
fig_biplot_enhanced = go.Figure()

# Add scatter plot for samples with cluster colors
cluster_colors = df['cluster_label'].values
unique_clusters = sorted(df['cluster_label'].dropna().unique())

for cluster_id in unique_clusters:
    cluster_mask = df['cluster_label'] == cluster_id
    cluster_samples = df[cluster_mask]
    
    fig_biplot_enhanced.add_scatter(
        x=pca_scores[cluster_mask, 0],
        y=pca_scores[cluster_mask, 1],
        mode='markers',
        marker=dict(
            size=10,
            color=DEFAULT_PLOTLY_COLORS[int(cluster_id) % len(DEFAULT_PLOTLY_COLORS)],
            line=dict(width=1, color='white'),
            opacity=0.8
        ),
        text=cluster_samples['Sample'],
        name=f'Cluster {int(cluster_id)}',
        hovertemplate='<b>%{text}</b><br>PC1: %{x:.3f}<br>PC2: %{y:.3f}<extra></extra>'
    )

# Add loading vectors with enhanced styling
scale_factor = 4.0  # Adjust for better visibility
arrow_colors = px.colors.qualitative.Set2

for i, feature in enumerate(input_columns):
    # Calculate arrow position
    arrow_x = loadings_biplot[i, 0] * scale_factor
    arrow_y = loadings_biplot[i, 1] * scale_factor
    
    # Add arrow as a line with annotation
    fig_biplot_enhanced.add_trace(go.Scatter(
        x=[0, arrow_x],
        y=[0, arrow_y],
        mode='lines',
        line=dict(
            color=arrow_colors[i % len(arrow_colors)],
            width=3
        ),
        showlegend=False,
        hoverinfo='skip'
    ))
    
    # Add arrowhead
    fig_biplot_enhanced.add_annotation(
        x=arrow_x,
        y=arrow_y,
        ax=0.0,
        ay=0.0,
        xref='x',
        yref='y',
        axref='x',
        ayref='y',
        showarrow=True,
        arrowhead=2,
        arrowsize=1.5,
        arrowwidth=2,
        #arrowcolor=arrow_colors[i % len(arrow_colors)],
        arrowcolor='darkblue',
    )
    
    # Add feature labels with better positioning
    label_offset = 1.05
    fig_biplot_enhanced.add_annotation(
        x=arrow_x * label_offset,
        y=arrow_y * label_offset,
        text=f'<b>{feature}</b>',
        showarrow=False,
        font=dict(
            size=16,
            color=arrow_colors[i % len(arrow_colors)],
            family='Arial Black'
        ),
        # bgcolor='rgba(255, 255, 255, 0.85)',
        # bordercolor=arrow_colors[i % len(arrow_colors)],
        # borderwidth=2,
        # borderpad=4
    )

# Add grid lines at origin
fig_biplot_enhanced.add_hline(y=0, line_dash="dash", line_color="gray", opacity=0.5)
fig_biplot_enhanced.add_vline(x=0, line_dash="dash", line_color="gray", opacity=0.5)

# Update layout with enhanced styling
fig_biplot_enhanced.update_layout(
    title=dict(
        text=f'<b>PCA Biplot - Samples and Feature Loadings</b><br>' +
             f'<sub>Variance Explained: PC1={pca_biplot.explained_variance_ratio_[0]*100:.1f}%, ' +
             f'PC2={pca_biplot.explained_variance_ratio_[1]*100:.1f}%</sub>',
        x=0.5,
        xanchor='center',
        font=dict(size=16)
    ),
    xaxis_title=f'<b>PC1 ({pca_biplot.explained_variance_ratio_[0]*100:.1f}%)</b>',
    yaxis_title=f'<b>PC2 ({pca_biplot.explained_variance_ratio_[1]*100:.1f}%)</b>',
    width=1000,
    height=800,
    showlegend=True,
    legend=dict(
        title=dict(text='<b>Clusters</b>'),
        orientation='v',
        x=1.02,
        y=1.0,
        bgcolor='rgba(255, 255, 255, 0.8)',
        bordercolor='black',
        borderwidth=1
    ),
    plot_bgcolor='rgba(245, 245, 245, 0.5)',
    xaxis=dict(
        gridcolor='white',
        gridwidth=2,
        zeroline=True,
        zerolinewidth=2,
        zerolinecolor='black'
    ),
    yaxis=dict(
        gridcolor='white',
        gridwidth=2,
        zeroline=True,
        zerolinewidth=2,
        zerolinecolor='black',
        scaleanchor='x',
        scaleratio=1
    ),
    font=dict(size=14)
)

# Save the biplot figure
fig_biplot_enhanced.write_html('PCA_Biplot_Enhanced.html')
print("Enhanced PCA biplot saved to 'PCA_Biplot_Enhanced.html'")


#show_figure(fig_biplot_enhanced)

# Print loading values in a formatted table
print("\nFeature Loadings on PC1 and PC2:")
print("-" * 50)
print(f"{'Feature':<12} {'PC1 Loading':>12} {'PC2 Loading':>12} {'Vector Length':>14}")
print("-" * 50)
for i, feature in enumerate(input_columns):
    vector_length = np.sqrt(loadings_biplot[i, 0]**2 + loadings_biplot[i, 1]**2)
    print(f"{feature:<12} {loadings_biplot[i, 0]:>12.4f} {loadings_biplot[i, 1]:>12.4f} {vector_length:>14.4f}")

print("="*50 + "\n")


# # Calculate PC2 isoline value (mean across all neurons)
# print("\n" + "="*50)
# print("PC2 Isoline Value Calculation")
# print("="*50)

# Calculate PC2 value for each neuron
pc2_neuron_map = np.zeros((final_nx, final_ny))

for i in range(final_nx):
    for j in range(final_ny):
        # Get weights for this neuron across all features
        neuron_weights = weights[i, j, :]
        
        # Calculate PC2 value as dot product of neuron weights and PC2 loadings
        pc2_neuron_map[i, j] = np.dot(neuron_weights, pca.components_[1, :])

# Use mean PC2 value as the isoline threshold
pc2_isoline_value = np.mean(pc2_neuron_map)

print(f"\nPC2 Isoline Value (mean): {pc2_isoline_value:.6f}")
print(f"PC2 range: [{pc2_neuron_map.min():.6f}, {pc2_neuron_map.max():.6f}]")
print("="*50 + "\n")


# # Print PC1 isoline value (red dotted line)
# print("\n" + "="*50)
# print("PC1 Isoline Value (Red Dotted Line)")
# print("="*50)
# print(f"PC1 Isoline Value: {pc1_isoline_value:.6f}")
# print(f"This value is constant across all SOM planes")
# print("="*50 + "\n")

# Add PC2 isolines to all feature plots
print("\n" + "="*50)
print("Adding PC2 Isolines to SOM Planes")
print("="*50)

def add_pc2_isolines(fig, coords, hexagons, feature_name, row, col):
    """
    Add PC2 isoline to a SOM plane based on PC2 neuron values.
    This ensures the same geometric line across all feature plots.
    """
    xx, yy = coords
    
    # Create interpolated grid for smoother contours using PC2 neuron map
    points = np.column_stack([xx.flatten(), yy.flatten()])
    values = pc2_neuron_map.flatten()
    
    # Create fine grid
    grid_x = np.linspace(xx.min(), xx.max(), 100)
    grid_y = np.linspace(yy.min(), yy.max(), 100)
    grid_xx, grid_yy = np.meshgrid(grid_x, grid_y)
    
    # Interpolate PC2 values
    grid_z = griddata(points, values, (grid_xx, grid_yy), method='cubic')
    
    # Generate contour for PC2 isoline value
    fig_temp = plt.figure()
    ax_temp = fig_temp.add_subplot(111)
    contours = ax_temp.contour(grid_xx, grid_yy, grid_z, levels=[pc2_isoline_value], colors='black', linewidths=2)
    plt.close(fig_temp)
    
    # Extract contour paths and add to plotly figure
    if len(contours.allsegs) > 0:
        for path in contours.allsegs[0]:
            fig.add_scatter(
                x=path[:, 0],
                y=path[:, 1],
                mode='lines',
                line=dict(color='darkblue', width=5.0, dash='dot'),
                showlegend=False,
                hoverinfo='skip',
                row=row,
                col=col
            )

# Add PC2 isolines to all feature plots (excluding pH_L if needed)
pc2_features = input_columns.copy()

for feat_name in pc2_features:
    if feat_name in input_columns:
        feat_idx = input_columns.index(feat_name)
        plot_position = feat_idx
        row_idx = subplot_ix[plot_position][0]
        col_idx = subplot_ix[plot_position][1]
        
        add_pc2_isolines(
            features_fig,
            (xx, yy),
            hexagons,
            feat_name,
            row_idx,
            col_idx
        )

print(f"PC2 isolines added to all feature plots")
print(f"PC2 Isoline Value: {pc2_isoline_value:.6f}")
print("="*50 + "\n")

# Also add PC2 isolines to the enhanced figures (with boundaries and with boundaries+isolines)
if 'fig_with_boundaries' in locals():
    for feat_name in pc2_features:
        if feat_name in input_columns:
            feat_idx = input_columns.index(feat_name)
            plot_position = feat_idx
            row_idx = subplot_ix_boundaries[plot_position][0]
            col_idx = subplot_ix_boundaries[plot_position][1]
            
            add_pc2_isolines(
                fig_with_boundaries,
                (xx, yy),
                hexagons,
                feat_name,
                row_idx,
                col_idx
            )
    print("PC2 isolines added to boundaries figure")

if 'fig_with_boundaries_isolines' in locals():
    for feat_name in pc2_features:
        if feat_name in input_columns:
            feat_idx = input_columns.index(feat_name)
            plot_position = feat_idx
            row_idx = subplot_ix_boundaries_isolines[plot_position][0]
            col_idx = subplot_ix_boundaries_isolines[plot_position][1]
            
            add_pc2_isolines(
                fig_with_boundaries_isolines,
                (xx, yy),
                hexagons,
                feat_name,
                row_idx,
                col_idx
            )
    print("PC2 isolines added to boundaries+isolines figure")

print("="*50 + "\n")

# Display updated figures
show_figure(features_fig)




# Create a single SOM plane with all isolines (excluding green dash isolines)
print("\n" + "="*50)
print("Creating Single SOM Plane with All Isolines")
print("="*50)

# Create a new figure for single plane with isolines
fig_single_isolines = go.Figure()

# Add hexagons with light background (using U-matrix or a neutral colorscale)
# Using a very light color scheme
for i in range(final_nx):
    for j in range(final_ny):
        fig_single_isolines.add_scatter(
            x=hexagons[i, j, :, 0],
            y=hexagons[i, j, :, 1],
            fill='toself',
            mode='lines',
            fillcolor='rgba(240, 240, 240, 0.5)',  # Light gray background
            line=dict(color='lightgray', width=0.5),
            showlegend=False,
            hoverinfo='skip',
        )

# Add cluster boundaries
for k, boundary in cluster_boundaries.items():
    if boundary.geom_type == 'Polygon':
        bnds = [boundary]
    elif boundary.geom_type == 'MultiPolygon':
        bnds = list(boundary.geoms)
    else:
        raise NotImplementedError(boundary.geom_type)
    
    for poly in bnds:
        if poly.boundary.geom_type == 'LineString':
            crds = np.array(poly.boundary.coords)
        else:
            crds = np.array(poly.boundary.geoms[0].coords)

        fig_single_isolines.add_scatter(
            x=crds[:, 0],
            y=crds[:, 1],
            mode="lines",
            showlegend=False,
            line=dict(color='black', width=2),
        )

# Add colored isolines (darkorange solid lines)
for feat_name, color_dict in isoline_colors.items():
    if feat_name in input_columns:
        feat_idx = input_columns.index(feat_name)
        som_weights = weights[:, :, feat_idx]
        
        points = np.column_stack([xx.flatten(), yy.flatten()])
        values = som_weights.flatten()
        
        grid_x = np.linspace(xx.min(), xx.max(), 100)
        grid_y = np.linspace(yy.min(), yy.max(), 100)
        grid_xx, grid_yy = np.meshgrid(grid_x, grid_y)
        
        grid_z = griddata(points, values, (grid_xx, grid_yy), method='cubic')
        
        for isoline_val, color in color_dict.items():
            fig_temp = plt.figure()
            ax_temp = fig_temp.add_subplot(111)
            contours = ax_temp.contour(grid_xx, grid_yy, grid_z, levels=[isoline_val], colors='black', linewidths=3.0)
            plt.close(fig_temp)
            
            line_style = 'solid'
            if feat_name in isoline_styles and isoline_val in isoline_styles[feat_name]:
                line_style = isoline_styles[feat_name][isoline_val]
            
            for level_idx in range(len(contours.levels)):
                paths = contours.allsegs[level_idx]
                for path in paths:
                    fig_single_isolines.add_scatter(
                        x=path[:, 0],
                        y=path[:, 1],
                        mode='lines',
                        line=dict(color=color, width=4.0, dash=line_style),
                        showlegend=False,
                        hoverinfo='skip',
                        name=f'{feat_name}={isoline_val:.2f}'
                    )

# Add PC1 isolines (red dotted)
points = np.column_stack([xx.flatten(), yy.flatten()])
values = pc1_neuron_map.flatten()

grid_x = np.linspace(xx.min(), xx.max(), 100)
grid_y = np.linspace(yy.min(), yy.max(), 100)
grid_xx, grid_yy = np.meshgrid(grid_x, grid_y)

grid_z = griddata(points, values, (grid_xx, grid_yy), method='cubic')

fig_temp = plt.figure()
ax_temp = fig_temp.add_subplot(111)
contours = ax_temp.contour(grid_xx, grid_yy, grid_z, levels=[pc1_isoline_value], colors='black', linewidths=2)
plt.close(fig_temp)

if len(contours.allsegs) > 0:
    for path in contours.allsegs[0]:
        fig_single_isolines.add_scatter(
            x=path[:, 0],
            y=path[:, 1],
            mode='lines',
            line=dict(color='red', width=4.5, dash='dot'),
            showlegend=False,
            hoverinfo='skip',
            name='PC1 Isoline'
        )

# Add PC2 isolines (blue dotted)
values = pc2_neuron_map.flatten()
grid_z = griddata(points, values, (grid_xx, grid_yy), method='cubic')

fig_temp = plt.figure()
ax_temp = fig_temp.add_subplot(111)
contours = ax_temp.contour(grid_xx, grid_yy, grid_z, levels=[pc2_isoline_value], colors='black', linewidths=2)
plt.close(fig_temp)

if len(contours.allsegs) > 0:
    for path in contours.allsegs[0]:
        fig_single_isolines.add_scatter(
            x=path[:, 0],
            y=path[:, 1],
            mode='lines',
            line=dict(color='darkblue', width=4.5, dash='dot'),
            showlegend=False,
            hoverinfo='skip',
            name='PC2 Isoline'
        )

# Update layout
fig_single_isolines.update_layout(
    title='SOM Plane with All Isolines (Excluding Green Dash Lines)',
    xaxis=dict(showticklabels=False, showgrid=False),
    yaxis=dict(showticklabels=False, showgrid=False, scaleanchor='x'),
    width=800,
    height=800,
    plot_bgcolor='white',
    showlegend=False
)

# Save and display
fig_single_isolines.write_html('SOM_Single_Plane_All_Isolines.html')
print("Single SOM plane with all isolines saved to 'SOM_Single_Plane_All_Isolines.html'")
#show_figure(fig_single_isolines)

print("="*50 + "\n")


# Create enhanced biplot with improved visualization
print("\n" + "="*50)
print("Creating Enhanced PCA Biplot")
print("="*50)

# Perform PCA on the transformed data (already done above, but ensure it's available)
pca_biplot = PCA(n_components=2)
pca_scores = pca_biplot.fit_transform(transformed_data)

# Get the loadings (principal components)
loadings_biplot = pca_biplot.components_.T * np.sqrt(pca_biplot.explained_variance_)

# Create the enhanced biplot
fig_biplot_enhanced = go.Figure()

# Add scatter plot for samples with cluster colors
cluster_colors = df['cluster_label'].values
unique_clusters = sorted(df['cluster_label'].dropna().unique())

for cluster_id in unique_clusters:
    cluster_mask = df['cluster_label'] == cluster_id
    cluster_samples = df[cluster_mask]
    
    fig_biplot_enhanced.add_scatter(
        x=pca_scores[cluster_mask, 0],
        y=pca_scores[cluster_mask, 1],
        mode='markers',
        marker=dict(
            size=10,
            color=DEFAULT_PLOTLY_COLORS[int(cluster_id) % len(DEFAULT_PLOTLY_COLORS)],
            line=dict(width=1, color='white'),
            opacity=0.8
        ),
        text=cluster_samples['Sample'],
        name=f'Cluster {int(cluster_id)}',
        hovertemplate='<b>%{text}</b><br>PC1: %{x:.3f}<br>PC2: %{y:.3f}<extra></extra>'
    )

# Add loading vectors with enhanced styling
scale_factor = 4.0  # Adjust for better visibility

# Mapping of original feature names to display names
feature_display_names = {
    'F4': 'f1',
    'F5': 'f2',
    'F7': 'f3',
    'F8': 'f4',
    'F9': 'f5'
}

for i, feature in enumerate(input_columns):
    # Calculate arrow position
    arrow_x = loadings_biplot[i, 0] * scale_factor
    arrow_y = loadings_biplot[i, 1] * scale_factor
    
    # Add arrow as a line with annotation
    fig_biplot_enhanced.add_trace(go.Scatter(
        x=[0, arrow_x],
        y=[0, arrow_y],
        mode='lines',
        line=dict(
            color='black',
            width=3
        ),
        showlegend=False,
        hoverinfo='skip'
    ))
    
    # Add arrowhead
    fig_biplot_enhanced.add_annotation(
        x=arrow_x,
        y=arrow_y,
        ax=0.0,
        ay=0.0,
        xref='x',
        yref='y',
        axref='x',
        ayref='y',
        showarrow=True,
        arrowhead=2,
        arrowsize=1.5,
        arrowwidth=2,
        arrowcolor='black',
    )
    
    # Add feature labels with better positioning
    label_offset = 1.05
    # Use display name if available, otherwise use original name
    display_name = feature_display_names.get(feature, feature)
    
    fig_biplot_enhanced.add_annotation(
        x=arrow_x * label_offset,
        y=arrow_y * label_offset,
        text=f'<b>{display_name}</b>',
        showarrow=False,
        font=dict(
            size=16,
            color='black',
            family='Arial Black'
        ),
    )

# Add grid lines at origin
fig_biplot_enhanced.add_hline(y=0, line_dash="dash", line_color="gray", opacity=0.5)
fig_biplot_enhanced.add_vline(x=0, line_dash="dash", line_color="gray", opacity=0.5)

# Update layout with enhanced styling
fig_biplot_enhanced.update_layout(
    title=dict(
        text=f'<b>PCA Biplot - Samples and Feature Loadings</b><br>' +
             f'<sub>Variance Explained: PC1={pca_biplot.explained_variance_ratio_[0]*100:.1f}%, ' +
             f'PC2={pca_biplot.explained_variance_ratio_[1]*100:.1f}%</sub>',
        x=0.5,
        xanchor='center',
        font=dict(size=16)
    ),
    xaxis_title=f'<b>PC1 ({pca_biplot.explained_variance_ratio_[0]*100:.1f}%)</b>',
    yaxis_title=f'<b>PC2 ({pca_biplot.explained_variance_ratio_[1]*100:.1f}%)</b>',
    width=1000,
    height=800,
    showlegend=True,
    legend=dict(
        title=dict(text='<b>Clusters</b>'),
        orientation='v',
        x=1.02,
        y=1.0,
        bgcolor='rgba(255, 255, 255, 0.8)',
        bordercolor='black',
        borderwidth=1
    ),
    plot_bgcolor='rgba(245, 245, 245, 0.5)',
    xaxis=dict(
        gridcolor='white',
        gridwidth=2,
        zeroline=True,
        zerolinewidth=2,
        zerolinecolor='black'
    ),
    yaxis=dict(
        gridcolor='white',
        gridwidth=2,
        zeroline=True,
        zerolinewidth=2,
        zerolinecolor='black',
        scaleanchor='x',
        scaleratio=1
    ),
    font=dict(size=14)
)

# Save the biplot figure
fig_biplot_enhanced.write_html('PCA_Biplot_Enhanced.html')
print("Enhanced PCA biplot saved to 'PCA_Biplot_Enhanced.html'")


#show_figure(fig_biplot_enhanced)

# Print loading values in a formatted table
print("\nFeature Loadings on PC1 and PC2:")
print("-" * 50)
print(f"{'Feature':<12} {'PC1 Loading':>12} {'PC2 Loading':>12} {'Vector Length':>14}")
print("-" * 50)
for i, feature in enumerate(input_columns):
    vector_length = np.sqrt(loadings_biplot[i, 0]**2 + loadings_biplot[i, 1]**2)
    print(f"{feature:<12} {loadings_biplot[i, 0]:>12.4f} {loadings_biplot[i, 1]:>12.4f} {vector_length:>14.4f}")

print("="*50 + "\n")

# Create enhanced biplot with improved visualization
print("\n" + "="*50)
print("Creating Enhanced PCA Biplot")
print("="*50)

# Perform PCA on the transformed data (already done above, but ensure it's available)
pca_biplot = PCA(n_components=2)
pca_scores = pca_biplot.fit_transform(transformed_data)

# Get the loadings (principal components)
loadings_biplot = pca_biplot.components_.T * np.sqrt(pca_biplot.explained_variance_)

# Create the enhanced biplot
fig_biplot_enhanced = go.Figure()

# Add scatter plot for samples with cluster colors
cluster_colors = df['cluster_label'].values
unique_clusters = sorted(df['cluster_label'].dropna().unique())

for cluster_id in unique_clusters:
    cluster_mask = df['cluster_label'] == cluster_id
    cluster_samples = df[cluster_mask]
    
    # Get base color and create darker border
    base_color = DEFAULT_PLOTLY_COLORS[int(cluster_id) % len(DEFAULT_PLOTLY_COLORS)]
    # Convert to darker shade for border
    import plotly.colors as pc
    import re
    
    # Parse RGB color string if it's in rgb() format, otherwise convert from hex
    if base_color.startswith('rgb'):
        rgb_match = re.findall(r'\d+', base_color)
        rgb = tuple(int(x) for x in rgb_match[:3])
    else:
        rgb = pc.hex_to_rgb(base_color)
    
    darker_rgb = tuple(int(c * 0.6) for c in rgb)  # 40% darker
    darker_color = f'rgb({darker_rgb[0]},{darker_rgb[1]},{darker_rgb[2]})'
    
    fig_biplot_enhanced.add_scatter(
        x=pca_scores[cluster_mask, 0],
        y=pca_scores[cluster_mask, 1],
        mode='markers',
        marker=dict(
            size=10,
            color=base_color,
            line=dict(width=2, color=darker_color),
            opacity=0.8
        ),
        text=cluster_samples['Sample'],
        name=f'Cluster {int(cluster_id)}',
        hovertemplate='<b>%{text}</b><br>PC1: %{x:.3f}<br>PC2: %{y:.3f}<extra></extra>'
    )

# Add loading vectors with enhanced styling
scale_factor = 4.0  # Adjust for better visibility

# Mapping of original feature names to display names
feature_display_names = {
    'F4': 'f1',
    'F5': 'f2',
    'F7': 'f3',
    'F8': 'f4',
    'F9': 'f5'
}

for i, feature in enumerate(input_columns):
    # Calculate arrow position
    arrow_x = loadings_biplot[i, 0] * scale_factor
    arrow_y = loadings_biplot[i, 1] * scale_factor
    
    # Add arrow as a line with annotation
    fig_biplot_enhanced.add_trace(go.Scatter(
        x=[0, arrow_x],
        y=[0, arrow_y],
        mode='lines',
        line=dict(
            color='black',
            width=1.5
        ),
        showlegend=False,
        hoverinfo='skip'
    ))
    
    # Add arrowhead
    fig_biplot_enhanced.add_annotation(
        x=arrow_x,
        y=arrow_y,
        ax=0.0,
        ay=0.0,
        xref='x',
        yref='y',
        axref='x',
        ayref='y',
        showarrow=True,
        arrowhead=2,
        arrowsize=1.5,
        arrowwidth=1.5,
        arrowcolor='black',
    )
    
    # Add feature labels with better positioning
    label_offset = 1.05
    # Use display name if available, otherwise use original name
    display_name = feature_display_names.get(feature, feature)
    
    fig_biplot_enhanced.add_annotation(
        x=arrow_x * label_offset,
        y=arrow_y * label_offset,
        text=f'<b>{display_name}</b>',
        showarrow=False,
        font=dict(
            size=16,
            color='black',
            family='Arial Black'
        ),
    )

# Add grid lines at origin
fig_biplot_enhanced.add_hline(y=0, line_dash="dash", line_color="gray", opacity=0.5)
fig_biplot_enhanced.add_vline(x=0, line_dash="dash", line_color="gray", opacity=0.5)

# Update layout with enhanced styling
fig_biplot_enhanced.update_layout(
    title=dict(
        text=f'<b>PCA Biplot - Samples and Feature Loadings</b><br>' +
             f'<sub>Variance Explained: PC1={pca_biplot.explained_variance_ratio_[0]*100:.1f}%, ' +
             f'PC2={pca_biplot.explained_variance_ratio_[1]*100:.1f}%</sub>',
        x=0.5,
        xanchor='center',
        font=dict(size=16)
    ),
    xaxis_title=f'<b>PC1 ({pca_biplot.explained_variance_ratio_[0]*100:.1f}%)</b>',
    yaxis_title=f'<b>PC2 ({pca_biplot.explained_variance_ratio_[1]*100:.1f}%)</b>',
    width=1000,
    height=800,
    showlegend=True,
    legend=dict(
        title=dict(text='<b>Clusters</b>'),
        orientation='v',
        x=1.02,
        y=1.0,
        bgcolor='rgba(255, 255, 255, 0.8)',
        bordercolor='black',
        borderwidth=1
    ),
    plot_bgcolor='rgba(245, 245, 245, 0.5)',
    xaxis=dict(
        gridcolor='white',
        gridwidth=2,
        zeroline=True,
        zerolinewidth=2,
        zerolinecolor='black',
        range=[-1, 1],
        tickfont=dict(size=16)
    ),
    yaxis=dict(
        gridcolor='white',
        gridwidth=2,
        zeroline=True,
        zerolinewidth=2,
        zerolinecolor='black',
        scaleanchor='x',
        scaleratio=1,
        range=[-1, 1],
        tickfont=dict(size=16)
    ),
    font=dict(size=14)
)

# Save the biplot figure
fig_biplot_enhanced.write_html('PCA_Biplot_Enhanced.html')
print("Enhanced PCA biplot saved to 'PCA_Biplot_Enhanced.html'")


#show_figure(fig_biplot_enhanced)

# Print loading values in a formatted table
print("\nFeature Loadings on PC1 and PC2:")
print("-" * 50)
print(f"{'Feature':<12} {'PC1 Loading':>12} {'PC2 Loading':>12} {'Vector Length':>14}")
print("-" * 50)
for i, feature in enumerate(input_columns):
    vector_length = np.sqrt(loadings_biplot[i, 0]**2 + loadings_biplot[i, 1]**2)
    print(f"{feature:<12} {loadings_biplot[i, 0]:>12.4f} {loadings_biplot[i, 1]:>12.4f} {vector_length:>14.4f}")

print("="*50 + "\n")


# Create enhanced biplot with improved visualization
print("\n" + "="*50)
print("Creating Enhanced PCA Biplot")
print("="*50)

# Perform PCA on the transformed data (already done above, but ensure it's available)
pca_biplot = PCA(n_components=2)
pca_scores = pca_biplot.fit_transform(transformed_data)

# Get the loadings (principal components)
loadings_biplot = pca_biplot.components_.T * np.sqrt(pca_biplot.explained_variance_)

# Create the enhanced biplot
fig_biplot_enhanced = go.Figure()

# Add scatter plot for samples with cluster colors
cluster_colors = df['cluster_label'].values
unique_clusters = sorted(df['cluster_label'].dropna().unique())

for cluster_id in unique_clusters:
    cluster_mask = df['cluster_label'] == cluster_id
    cluster_samples = df[cluster_mask]
    
    # Get base color and create darker border
    base_color = DEFAULT_PLOTLY_COLORS[int(cluster_id) % len(DEFAULT_PLOTLY_COLORS)]
    # Convert to darker shade for border
    import plotly.colors as pc
    
    # Parse RGB color string if it's in rgb() format, otherwise convert from hex
    if base_color.startswith('rgb'):
        rgb_match = re.findall(r'\d+', base_color)
        rgb = tuple(int(x) for x in rgb_match[:3])
    else:
        rgb = pc.hex_to_rgb(base_color)
    
    darker_rgb = tuple(int(c * 0.6) for c in rgb)  # 40% darker
    darker_color = f'rgb({darker_rgb[0]},{darker_rgb[1]},{darker_rgb[2]})'
    
    fig_biplot_enhanced.add_scatter(
        x=pca_scores[cluster_mask, 0],
        y=pca_scores[cluster_mask, 1],
        mode='markers',
        marker=dict(
            size=10,
            color=base_color,
            line=dict(width=2, color=darker_color),
            opacity=0.8
        ),
        text=cluster_samples['Sample'],
        name=f'Cluster {int(cluster_id)}',
        hovertemplate='<b>%{text}</b><br>PC1: %{x:.3f}<br>PC2: %{y:.3f}<extra></extra>'
    )

# Add loading vectors with enhanced styling
scale_factor = 4.0  # Adjust for better visibility

# Mapping of original feature names to display names
feature_display_names = {
    'F4': 'f1',
    'F5': 'f2',
    'F7': 'f3',
    'F8': 'f4',
    'F9': 'f5'
}

for i, feature in enumerate(input_columns):
    # Calculate arrow position
    arrow_x = loadings_biplot[i, 0] * scale_factor
    arrow_y = loadings_biplot[i, 1] * scale_factor
    
    # Add arrow as a line with annotation
    fig_biplot_enhanced.add_trace(go.Scatter(
        x=[0, arrow_x],
        y=[0, arrow_y],
        mode='lines',
        line=dict(
            color='black',
            width=1.5
        ),
        showlegend=False,
        hoverinfo='skip'
    ))
    
    # Add arrowhead
    fig_biplot_enhanced.add_annotation(
        x=arrow_x,
        y=arrow_y,
        ax=0.0,
        ay=0.0,
        xref='x',
        yref='y',
        axref='x',
        ayref='y',
        showarrow=True,
        arrowhead=2,
        arrowsize=1.5,
        arrowwidth=1.5,
        arrowcolor='black',
    )
    
    # Add feature labels with better positioning
    label_offset = 1.05
    # Use display name if available, otherwise use original name
    display_name = feature_display_names.get(feature, feature)
    
    fig_biplot_enhanced.add_annotation(
        x=arrow_x * label_offset,
        y=arrow_y * label_offset,
        text=f'<b>{display_name}</b>',
        showarrow=False,
        font=dict(
            size=16,
            color='black',
            family='Arial Black'
        ),
    )

# Add grid lines at origin
fig_biplot_enhanced.add_hline(y=0, line_dash="dash", line_color="gray", opacity=0.5)
fig_biplot_enhanced.add_vline(x=0, line_dash="dash", line_color="gray", opacity=0.5)

# Add square frame around the plot
fig_biplot_enhanced.add_shape(
    type="rect",
    x0=-1, y0=-1, x1=1, y1=1,
    line=dict(color="black", width=2),
    fillcolor="rgba(0,0,0,0)"
)

# Update layout with enhanced styling
fig_biplot_enhanced.update_layout(
    title=dict(
        text=f'<b>PCA Biplot - Samples and Feature Loadings</b><br>' +
             f'<sub>Variance Explained: PC1={pca_biplot.explained_variance_ratio_[0]*100:.1f}%, ' +
             f'PC2={pca_biplot.explained_variance_ratio_[1]*100:.1f}%</sub>',
        x=0.5,
        xanchor='center',
        font=dict(size=16)
    ),
    xaxis_title=f'<b>PC1 ({pca_biplot.explained_variance_ratio_[0]*100:.1f}%)</b>',
    yaxis_title=f'<b>PC2 ({pca_biplot.explained_variance_ratio_[1]*100:.1f}%)</b>',
    width=1000,
    height=800,
    showlegend=True,
    legend=dict(
        title=dict(text='<b>Clusters</b>'),
        orientation='v',
        x=1.02,
        y=1.0,
        bgcolor='rgba(255, 255, 255, 0.8)',
        bordercolor='black',
        borderwidth=1
    ),
    plot_bgcolor='rgba(245, 245, 245, 0.5)',
    xaxis=dict(
        gridcolor='white',
        gridwidth=2,
        zeroline=True,
        zerolinewidth=2,
        zerolinecolor='black',
        range=[-1, 1],
        tickfont=dict(size=16),
        constrain='domain'
    ),
    yaxis=dict(
        gridcolor='white',
        gridwidth=2,
        zeroline=True,
        zerolinewidth=2,
        zerolinecolor='black',
        scaleanchor='x',
        scaleratio=1,
        range=[-1, 1],
        tickfont=dict(size=18),
        constrain='domain'
    ),
    font=dict(size=16)
)

# Save the biplot figure
fig_biplot_enhanced.write_html('PCA_Biplot_Enhanced.html')
print("Enhanced PCA biplot saved to 'PCA_Biplot_Enhanced.html'")


show_figure(fig_biplot_enhanced)

# Print loading values in a formatted table
print("\nFeature Loadings on PC1 and PC2:")
print("-" * 50)
print(f"{'Feature':<12} {'PC1 Loading':>12} {'PC2 Loading':>12} {'Vector Length':>14}")
print("-" * 50)
for i, feature in enumerate(input_columns):
    vector_length = np.sqrt(loadings_biplot[i, 0]**2 + loadings_biplot[i, 1]**2)
    print(f"{feature:<12} {loadings_biplot[i, 0]:>12.4f} {loadings_biplot[i, 1]:>12.4f} {vector_length:>14.4f}")

print("="*50 + "\n")
# Create enhanced biplot with improved visualization
print("\n" + "="*50)
print("Creating Enhanced PCA Biplot")
print("="*50)

# Perform PCA on the transformed data (already done above, but ensure it's available)
pca_biplot = PCA(n_components=2)
pca_scores = pca_biplot.fit_transform(transformed_data)

# Get the loadings (principal components)
loadings_biplot = pca_biplot.components_.T * np.sqrt(pca_biplot.explained_variance_)

# Create the enhanced biplot
fig_biplot_enhanced = go.Figure()

# Add scatter plot for samples with cluster colors
cluster_colors = df['cluster_label'].values
unique_clusters = sorted(df['cluster_label'].dropna().unique())

for cluster_id in unique_clusters:
    cluster_mask = df['cluster_label'] == cluster_id
    cluster_samples = df[cluster_mask]
    
    # Get base color and create darker border
    base_color = DEFAULT_PLOTLY_COLORS[int(cluster_id) % len(DEFAULT_PLOTLY_COLORS)]
    # Convert to darker shade for border
    import plotly.colors as pc
    
    # Parse RGB color string if it's in rgb() format, otherwise convert from hex
    if base_color.startswith('rgb'):
        rgb_match = re.findall(r'\d+', base_color)
        rgb = tuple(int(x) for x in rgb_match[:3])
    else:
        rgb = pc.hex_to_rgb(base_color)
    
    darker_rgb = tuple(int(c * 0.6) for c in rgb)  # 40% darker
    darker_color = f'rgb({darker_rgb[0]},{darker_rgb[1]},{darker_rgb[2]})'
    
    fig_biplot_enhanced.add_scatter(
        x=pca_scores[cluster_mask, 0],
        y=pca_scores[cluster_mask, 1],
        mode='markers',
        marker=dict(
            size=10,
            color=base_color,
            line=dict(width=2, color=darker_color),
            opacity=0.8
        ),
        text=cluster_samples['Sample'],
        name=f'Cluster {int(cluster_id)}',
        hovertemplate='<b>%{text}</b><br>PC1: %{x:.3f}<br>PC2: %{y:.3f}<extra></extra>'
    )

# Add loading vectors with enhanced styling
scale_factor = 4.0  # Adjust for better visibility

# Mapping of original feature names to display names
feature_display_names = {
    'F4': 'f1',
    'F5': 'f2',
    'F7': 'f3',
    'F8': 'f4',
    'F9': 'f5'
}

for i, feature in enumerate(input_columns):
    # Calculate arrow position
    arrow_x = loadings_biplot[i, 0] * scale_factor
    arrow_y = loadings_biplot[i, 1] * scale_factor
    
    # Add arrow as a line with annotation
    fig_biplot_enhanced.add_trace(go.Scatter(
        x=[0, arrow_x],
        y=[0, arrow_y],
        mode='lines',
        line=dict(
            color='black',
            width=1.5
        ),
        showlegend=False,
        hoverinfo='skip'
    ))
    
    # Add arrowhead
    fig_biplot_enhanced.add_annotation(
        x=arrow_x,
        y=arrow_y,
        ax=0.0,
        ay=0.0,
        xref='x',
        yref='y',
        axref='x',
        ayref='y',
        showarrow=True,
        arrowhead=2,
        arrowsize=1.5,
        arrowwidth=1.5,
        arrowcolor='black',
    )
    
    # Add feature labels with better positioning
    label_offset = 1.05
    # Use display name if available, otherwise use original name
    display_name = feature_display_names.get(feature, feature)
    
    fig_biplot_enhanced.add_annotation(
        x=arrow_x * label_offset,
        y=arrow_y * label_offset,
        text=f'<b>{display_name}</b>',
        showarrow=False,
        font=dict(
            size=16,
            color='black',
            family='Arial Black'
        ),
    )

# Add grid lines at origin
fig_biplot_enhanced.add_hline(y=0, line_dash="dash", line_color="gray", opacity=0.5)
fig_biplot_enhanced.add_vline(x=0, line_dash="dash", line_color="gray", opacity=0.5)

# Add square frame around the plot
fig_biplot_enhanced.add_shape(
    type="rect",
    x0=-1, y0=-1, x1=1, y1=1,
    line=dict(color="black", width=2),
    fillcolor="rgba(0,0,0,0)"
)

# Update layout with enhanced styling
fig_biplot_enhanced.update_layout(
    title=dict(
        text=f'<b>PCA Biplot - Samples and Feature Loadings</b><br>' +
             f'<sub>Variance Explained: PC1={pca_biplot.explained_variance_ratio_[0]*100:.1f}%, ' +
             f'PC2={pca_biplot.explained_variance_ratio_[1]*100:.1f}%</sub>',
        x=0.5,
        xanchor='center',
        font=dict(size=16)
    ),
    xaxis_title=f'<b>PC1 ({pca_biplot.explained_variance_ratio_[0]*100:.1f}%)</b>',
    yaxis_title=f'<b>PC2 ({pca_biplot.explained_variance_ratio_[1]*100:.1f}%)</b>',
    width=1000,
    height=800,
    showlegend=True,
    legend=dict(
        title=dict(text='<b>Clusters</b>'),
        orientation='v',
        x=1.02,
        y=1.0,
        bgcolor='rgba(255, 255, 255, 0.8)',
        bordercolor='black',
        borderwidth=1
    ),
    plot_bgcolor='rgba(245, 245, 245, 0.5)',
    xaxis=dict(
        gridcolor='white',
        gridwidth=2,
        zeroline=True,
        zerolinewidth=2,
        zerolinecolor='black',
        range=[-1, 1],
        tickfont=dict(size=16),
        constrain='domain',
        tickvals=[-1.0, -0.5, 0.0, 0.5, 1.0],
        ticktext=["-1.0", "-0.5", "0.0", "0.5", "1.0"]
    ),
    yaxis=dict(
        gridcolor='white',
        gridwidth=2,
        zeroline=True,
        zerolinewidth=2,
        zerolinecolor='black',
        scaleanchor='x',
        scaleratio=1,
        range=[-1, 1],
        tickfont=dict(size=18),
        constrain='domain',
        tickvals=[-1.0, -0.5, 0.0, 0.5, 1.0],
        ticktext=["-1.0", "-0.5", "0.0", "0.5", "1.0"]
    ),
    font=dict(size=16)
)

# Save the biplot figure
fig_biplot_enhanced.write_html('PCA_Biplot_Enhanced.html')
print("Enhanced PCA biplot saved to 'PCA_Biplot_Enhanced.html'")


show_figure(fig_biplot_enhanced)

# Print loading values in a formatted table
print("\nFeature Loadings on PC1 and PC2:")
print("-" * 50)
print(f"{'Feature':<12} {'PC1 Loading':>12} {'PC2 Loading':>12} {'Vector Length':>14}")
print("-" * 50)
for i, feature in enumerate(input_columns):
    vector_length = np.sqrt(loadings_biplot[i, 0]**2 + loadings_biplot[i, 1]**2)
    print(f"{feature:<12} {loadings_biplot[i, 0]:>12.4f} {loadings_biplot[i, 1]:>12.4f} {vector_length:>14.4f}")

print("="*50 + "\n")





